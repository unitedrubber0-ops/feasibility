# main.py - FINAL PRODUCTION VERSION with Backend Rendering

import os
import io
import json
import time
import PyPDF2
import docx
import fitz  # PyMuPDF for all document processing
from PIL import Image  # For GD&T image processing
import google.generativeai as genai
import google.api_core.exceptions
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS

# --- Helper function to extract text page by page ---
def extract_text_from_pdf_paginated(file_stream):
    try:
        pdf_reader = PyPDF2.PdfReader(file_stream)
        pages_text = [page.extract_text() or "" for page in pdf_reader.pages]
        return pages_text
    except Exception as e:
        print(f"Error reading PDF paginated: {e}")
        return None

# --- Existing helper functions ---
def extract_text_from_docx(file_stream):
    try:
        document = docx.Document(file_stream)
        return "\n".join([para.text for para in document.paragraphs])
    except Exception as e:
        print(f"Error reading DOCX: {e}")
        return None

def generate_with_retry(model, prompt, max_retries=3, retry_delay=5):
    for attempt in range(max_retries):
        try:
            response = model.generate_content(prompt)
            cleaned_text = response.text.strip().replace("```json", "").replace("```", "")
            json.loads(cleaned_text) # Validate JSON before returning
            return cleaned_text
        except Exception as e:
            print(f"Error on attempt {attempt + 1}: {e}. Retrying...")
            if attempt < max_retries - 1:
                time.sleep(retry_delay)
            else:
                raise e

# --- Flask App Initialization ---
app = Flask(__name__, static_url_path='')

@app.route('/')
def serve_index():
    return send_file('index.html')
CORS(app, 
    resources={
        r"/*": {
            "origins": [
                "https://feasibility-638d.onrender.com",  # Backend URL
                "https://feasibility-1.onrender.com",  # Frontend URL
                "http://127.0.0.1:5001",  # Local development URL
                "http://localhost:5001"  # Alternative local development URL
            ],
            "methods": ["GET", "POST", "OPTIONS"],
            "allow_headers": ["Content-Type", "Authorization"],
            "expose_headers": ["Content-Range", "X-Content-Range"],
            "supports_credentials": True
        }
    })

# --- Gemini API Configuration ---
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")
if GEMINI_API_KEY:
    genai.configure(api_key=GEMINI_API_KEY)

# --- Main API Endpoint ---
@app.route('/generate-report', methods=['POST'])
def generate_report_handler():
    print("Received request to /generate-report")
    if 'sourceFile' not in request.files:
        return jsonify({"error": "Source file is missing"}), 400

    source_file = request.files['sourceFile']

    try:
        source_stream = io.BytesIO(source_file.read())
        source_pages = []
        if source_file.filename.endswith('.pdf'):
            source_pages = extract_text_from_pdf_paginated(source_stream)
        elif source_file.filename.endswith('.docx'):
            source_pages.append(extract_text_from_docx(source_stream))
        else:
            source_pages.append(source_stream.read().decode('utf-8', errors='ignore'))
        
        if not source_pages:
            return jsonify({"error": "Could not extract text from the source file."}), 500
    except Exception as e:
        print(f"Error during file processing: {e}")
        return jsonify({"error": "Failed to process files"}), 500

    try:
        if not GEMINI_API_KEY:
             return jsonify({"error": "Gemini API key is not configured."}), 500

        safety_settings = [{"category": c, "threshold": "BLOCK_NONE"} for c in ["HARM_CATEGORY_HARASSMENT", "HARM_CATEGORY_HATE_SPEECH", "HARM_CATEGORY_SEXUALLY_EXPLICIT", "HARM_CATEGORY_DANGEROUS_CONTENT"]]
        generation_config = genai.GenerationConfig(max_output_tokens=8192, temperature=0.1)
        model = genai.GenerativeModel("gemini-2.5-pro", generation_config=generation_config, safety_settings=safety_settings)

        aggregated_header = {}
        aggregated_rows = []
        table_columns = None

        # --- FINAL, PRODUCTION-HARDENED PROMPT ---
        prompt_extract_template = """
        You are a highly robust data extraction assistant. Your task is to analyze potentially messy OCR text from a DOCUMENT PAGE and extract its header and table data into a perfect JSON format.

        **DOCUMENT PAGE TEXT:**
        ---
        {page_text}
        ---

        **CRITICAL INSTRUCTIONS:**
        1.  Your primary goal is to return a **complete and valid JSON object**. Do not stop halfway.
        2.  The JSON must have two top-level keys: "header" (an object) and "table" (an object).
        3.  The "table" object must contain "columns" (a list of strings) and "rows" (a list of lists of strings).
        4.  **Handling Multi-line Table Rows:** Some rows in the source text span multiple lines (e.g., a "Description" parameter). You MUST consolidate all parts of a single logical row into one list in the JSON "rows" array.
        5.  If no table exists on the page, the "rows" array must be an empty list `[]`.
        6.  If no header data exists, the "header" object must be an empty object `{}`.
        7.  Do not include this instructional text in your response. Your response must only be the raw JSON object.
        """

        for i, page_text in enumerate(source_pages):
            print(f"Executing AI on page {i + 1}/{len(source_pages)}...")
            if not page_text.strip():
                print(f"Skipping empty page {i + 1}.")
                continue
            
            prompt = prompt_extract_template.format(page_text=page_text)
            extracted_json_string = generate_with_retry(model, prompt)
            page_json = json.loads(extracted_json_string)

            if page_json.get("header"):
                aggregated_header.update(page_json["header"])
            
            if page_json.get("table") and page_json["table"].get("rows"):
                if table_columns is None and page_json["table"].get("columns"):
                    table_columns = page_json["table"]["columns"]
                aggregated_rows.extend(page_json["table"]["rows"])
        
        if table_columns is None: table_columns = []
        
        final_report_json = {"header": aggregated_header, "table": {"columns": table_columns, "rows": aggregated_rows}}
        print("AI extraction complete for all pages. Sending final report.")

        return jsonify(final_report_json)

    except Exception as e:
        print(f"An error occurred during the AI process: {e}")
        return jsonify({"error": "Failed to generate report from AI model."}), 500

@app.route('/analyze-gdt-at-point', methods=['POST'])
def analyze_gdt_at_point_handler():
    """
    Analyzes a cropped section of a drawing page using a compartment-based
    prompt with self-correction for the highest possible GD&T accuracy.
    """
    if 'sourceFile' not in request.files or 'x' not in request.form or 'y' not in request.form or 'page_num' not in request.form:
        return jsonify({"error": "Missing source file or coordinate data"}), 400

    source_file = request.files['sourceFile']
    x_coord = float(request.form['x'])
    y_coord = float(request.form['y'])
    page_num = int(request.form['page_num'])

    try:
        file_stream = io.BytesIO(source_file.read())
        file_type = source_file.filename.split('.')[-1]
        
        doc = fitz.open(stream=file_stream, filetype=file_type)
        if page_num < 1 or page_num > doc.page_count:
            return jsonify({"error": "Invalid page number"}), 400
            
        page = doc.load_page(page_num - 1)
        
        # Backend cropping logic
        CROP_WIDTH = 400  # A good size for context
        CROP_HEIGHT = 150
        clip_box = fitz.Rect(x_coord - CROP_WIDTH / 2, y_coord - CROP_HEIGHT / 2, x_coord + CROP_WIDTH / 2, y_coord + CROP_HEIGHT / 2)
        pix = page.get_pixmap(dpi=200, clip=clip_box)
        
        img_buffer = io.BytesIO()
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        img.save(img_buffer, format="PNG")
        img_buffer.seek(0)
        
        doc.close()
        
        model = genai.GenerativeModel("gemini-1.5-pro-latest")
        gdt_image = {'mime_type': 'image/png', 'data': img_buffer.getvalue()}

        # --- THE ULTIMATE PROMPT with SELF-CORRECTION ---
        prompt = [
            "You are a world-class expert in Geometric Dimensioning and Tolerancing (GD&T) following the ASME Y14.5 standard. Your task is to parse a cropped image of a Feature Control Frame with absolute precision.",
            "You must analyze the frame compartment by compartment from left to right.",
            
            "**PARSING RULES:**",
            "1. **First Compartment:** Identify the geometric characteristic symbol.",
            "2. **Second Compartment (Tolerance):** Extract the full tolerance value. Identify if a 'Ø' (diameter) symbol is present and if a material condition modifier 'Ⓜ' (MMC) or 'Ⓛ' (LMC) is present.",
            "3. **Third and Subsequent Compartments (Datums):** Identify the primary, secondary, and tertiary datums. For each datum, identify its own material condition modifier.",
            
            # --- NEW SELF-CORRECTION INSTRUCTION ---
            "CRITICAL: Double-check the numerical values. Tolerance values in this context are rarely small decimals like '0.9' when the number on the drawing is clearly '9'. Be careful to distinguish between periods and pixel noise.",

            "**EXAMPLE:** For an image showing `Position | Ø9 M | A M - B M | C`:",
            """
            {
              "gdt_symbol_name": "Position",
              "tolerance_value": "9",
              "diameter_symbol": true,
              "material_condition_modifier": "MMC",
              "datums": [
                { "datum_letter": "A", "datum_material_condition": "MMC" },
                { "datum_letter": "B", "datum_material_condition": "MMC" },
                { "datum_letter": "C", "datum_material_condition": null }
              ]
            }
            """,
            "CRITICAL: Do not infer any information from text outside the main rectangular frame.",
            "Now, analyze this image with extreme precision:",
            gdt_image
        ]
        
        response = model.generate_content(prompt)
        cleaned_text = response.text.strip().replace("```json", "").replace("```", "")
        response_json = json.loads(cleaned_text)
        
        return jsonify(response_json)

    except Exception as e:
        print(f"An error occurred during final GD&T analysis: {e}")
        return jsonify({"error": f"Failed to analyze GD&T feature: {str(e)}"}), 500

# --- The /export-docx endpoint remains unchanged ---
@app.route('/export-docx', methods=['POST'])
def export_docx_handler():
    # ... This function is correct and does not need to be changed ...
    data = request.get_json()
    try:
        document = docx.Document()
        document.add_heading('Inspection Report', level=1)
        if data.get("header"):
            for key, value in data.get('header', {}).items():
                document.add_paragraph(f"{key}: {value}")
        document.add_paragraph()
        if data.get("table"):
            table_data = data.get('table', {})
            columns = table_data.get('columns', [])
            rows = table_data.get('rows', [])
            if columns and rows:
                table = document.add_table(rows=1, cols=len(columns))
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                for i, col_name in enumerate(columns):
                    hdr_cells[i].text = col_name
                for row_data in rows:
                    row_cells = table.add_row().cells
                    for i, cell_text in enumerate(row_data):
                        row_cells[i].text = str(cell_text)
        file_stream = io.BytesIO()
        document.save(file_stream)
        file_stream.seek(0)
        return send_file(
            file_stream,
            as_attachment=True,
            download_name='inspection_report.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        print(f"Error creating DOCX file: {e}")
        return jsonify({"error": "Failed to create DOCX file"}), 500

@app.route('/process-document-for-ocr', methods=['POST'])
def process_document_for_ocr_handler():
    """
    Performs OCR using PyMuPDF and returns coordinate data AND page count.
    """
    if 'sourceFile' not in request.files:
        return jsonify({"error": "Missing source file"}), 400

    source_file = request.files['sourceFile']
    
    try:
        file_stream = io.BytesIO(source_file.read())
        file_type = source_file.filename.split('.')[-1]
        doc = fitz.open(stream=file_stream, filetype=file_type)
        
        ocr_results = []
        for page_num, page in enumerate(doc):
            words = page.get_text("words")
            page_data = {
                "page": page_num + 1,
                "width": page.rect.width, "height": page.rect.height,
                "words": [{"text": w[4], "bbox": [w[0], w[1], w[2], w[3]]} for w in words]
            }
            ocr_results.append(page_data)
        
        # Return both the page count and the OCR results
        response_data = {"page_count": doc.page_count, "ocr_results": ocr_results}
        doc.close()
        return jsonify(response_data)

    except Exception as e:
        print(f"An error occurred during OCR processing: {e}")
        return jsonify({"error": f"Failed to process document: {str(e)}"}), 500

@app.route('/get-page-as-image/<int:page_num>', methods=['POST'])
def get_page_as_image_handler(page_num):
    """
    Renders a specific page of a PDF or DOCX file as a PNG image.
    """
    if 'sourceFile' not in request.files:
        return jsonify({"error": "Missing source file"}), 400
        
    source_file = request.files['sourceFile']
    try:
        file_stream = io.BytesIO(source_file.read())
        file_type = source_file.filename.split('.')[-1]
        doc = fitz.open(stream=file_stream, filetype=file_type)

        if page_num < 1 or page_num > doc.page_count:
            return jsonify({"error": "Invalid page number"}), 400

        page = doc.load_page(page_num - 1) # Page numbers are 0-indexed in PyMuPDF
        pix = page.get_pixmap(dpi=150) # Render at 150 DPI for good quality
        img_byte_arr = io.BytesIO(pix.tobytes("png"))
        doc.close()

        return send_file(img_byte_arr, mimetype='image/png')

    except Exception as e:
        print(f"Error rendering page as image: {e}")
        return jsonify({"error": f"Failed to render page: {str(e)}"}), 500


@app.route('/get-value-for-label', methods=['POST'])
def get_value_for_label_handler():
    """
    Receives a document and a specific label, and returns the value for that label.
    This is a fast, targeted endpoint for the ballooning feature.
    """
    if 'sourceFile' not in request.files or 'label' not in request.form:
        return jsonify({"error": "Missing source file or label"}), 400

    source_file = request.files['sourceFile']
    label = request.form['label']

    try:
        source_stream = io.BytesIO(source_file.read())
        doc = fitz.open(stream=source_stream, filetype=source_file.filename.split('.')[-1])
        source_text = "".join([page.get_text() for page in doc])
        doc.close()

        if not source_text:
            return jsonify({"error": "Could not extract text from source file."}), 500

    except Exception as e:
        return jsonify({"error": f"Failed to process file: {e}"}), 500

    try:
        # --- THIS IS THE FIX ---
        # Using the latest, stable model name for Gemini Flash.
        model = genai.GenerativeModel("gemini-1.5-pro-latest")

        prompt = f"""
        You are a data extraction specialist. In the following DOCUMENT TEXT, find the label "{label}" and return its corresponding value.

        **DOCUMENT TEXT:**
        ---
        {source_text}
        ---

        **INSTRUCTIONS:**
        Return a single, raw JSON object with two keys: "parameter" and "value".
        - "parameter" should be the exact label that was found (e.g., "HOSE_ID").
        - "value" should be the numerical or text value associated with that label (e.g., "24.6").
        """

        # We can reuse the robust retry function you already have
        response_text = generate_with_retry(model, prompt)
        response_json = json.loads(response_text)
        
        return jsonify(response_json)

    except Exception as e:
        print(f"An error occurred during AI processing for label '{label}': {e}")
        return jsonify({"error": f"AI processing failed: {str(e)}"}), 500

# Debug route to verify CORS
@app.route('/debug-cors', methods=['GET', 'OPTIONS'])
def debug_cors():
    """Simple endpoint to verify CORS configuration."""
    return jsonify({
        "message": "CORS is working",
        "origin": request.headers.get('Origin', 'No origin header'),
        "method": request.method
    })

def log_request_info():
    """Log detailed information about the current request."""
    print("\n=== Request Information ===")
    print(f"Method: {request.method}")
    print(f"URL: {request.url}")
    print(f"Headers: {dict(request.headers)}")
    print(f"Files: {list(request.files.keys()) if request.files else 'No files'}")
    print(f"Form Data: {dict(request.form)}")
    print("=========================\n")

@app.before_request
def before_request():
    """Log information about each request before processing."""
    log_request_info()

@app.after_request
def after_request(response):
    """Ensure CORS headers are present on all responses."""
    print(f"\n=== Response Information ===")
    print(f"Status: {response.status_code}")
    print(f"Headers: {dict(response.headers)}")
    print("=========================\n")
    
    # Always add CORS headers
    if request.headers.get('Origin'):
        response.headers['Access-Control-Allow-Origin'] = request.headers['Origin']
        response.headers['Access-Control-Allow-Credentials'] = 'true'
        response.headers['Access-Control-Allow-Methods'] = 'GET, POST, OPTIONS'
        response.headers['Access-Control-Allow-Headers'] = 'Content-Type'
    
    return response

@app.errorhandler(Exception)
def handle_exception(e):
    """Log any uncaught exceptions and ensure proper error response."""
    print(f"\n!!! UNCAUGHT EXCEPTION !!!")
    print(f"Type: {type(e).__name__}")
    print(f"Message: {str(e)}")
    
    import traceback
    print("\nTraceback:")
    print(traceback.format_exc())
    
    import sys
    print("\nSystem Info:")
    print(f"Python version: {sys.version}")
    print(f"Platform: {sys.platform}")
    
    try:
        # Try to get memory info on Linux systems
        import resource
        print("\nMemory Usage:")
        print(f"Max RSS: {resource.getrusage(resource.RUSAGE_SELF).ru_maxrss} KB")
    except ImportError:
        print("Resource module not available")
    
    print("\n=========================\n")
    
    # Create the error response
    response = jsonify({
        "error": "Internal server error",
        "message": str(e),
        "type": type(e).__name__
    })
    response.status_code = 500
    
    # Add CORS headers
    if request.headers.get('Origin'):
        response.headers['Access-Control-Allow-Origin'] = request.headers['Origin']
        response.headers['Access-Control-Allow-Credentials'] = 'true'
        response.headers['Access-Control-Allow-Methods'] = 'GET, POST, OPTIONS'
        response.headers['Access-Control-Allow-Headers'] = 'Content-Type'
    
    return response

if __name__ == '__main__':
    app.run(debug=True, port=5001)